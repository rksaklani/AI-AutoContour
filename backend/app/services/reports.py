"""Report generation: HTML (Jinja2) -> PDF (WeasyPrint) + DOCX (python-docx).

Heavy/optional renderers (WeasyPrint, python-docx) are imported lazily so the module
always imports cleanly; if a renderer is unavailable that format is skipped gracefully.
"""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path

from jinja2 import Environment, FileSystemLoader, select_autoescape

from app.ai import get_engine
from app.core.config import settings
from app.core.logging import get_logger
from app.models.study import Study
from app.services import storage

logger = get_logger(__name__)

_TEMPLATE_DIR = Path(__file__).parent / "templates"
_env = Environment(
    loader=FileSystemLoader(str(_TEMPLATE_DIR)),
    autoescape=select_autoescape(["html", "xml"]),
)

_IMAGE_MODALITIES = {"CT", "MR", "MRI", "PT", "PET", "US", "CR", "DX", "NM"}


def _series_instances_for_study(study: Study) -> list[dict]:
    out: list[dict] = []
    for series in study.series:
        mod = (series.modality or "").upper()
        if mod not in _IMAGE_MODALITIES:
            continue
        for inst in sorted(series.instances, key=lambda i: (i.instance_number or 0)):
            out.append(
                {
                    "series_id": str(series.id),
                    "instance_id": str(inst.id),
                    "modality": mod,
                    "object_key": inst.object_key,
                    "url": storage.presigned_get_url(inst.object_key),
                    "instance_number": inst.instance_number,
                }
            )
    return out


def _context(study: Study) -> dict:
    ctx = {
        "app_name": settings.LUMIRA_APP_NAME,
        "engine_name": get_engine().name,
        "generated_at": datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC"),
        "study": study,
        "findings": sorted(study.findings, key=lambda f: f.confidence, reverse=True),
        "segmentations": study.segmentations,
        "ai_narrative": None,
        "ai_impression": None,
    }
    from app.ai import get_vila_engine

    vila = get_vila_engine()
    if vila is not None:
        findings_payload = [
            {
                "label": f.label,
                "location": f.location,
                "confidence": f.confidence,
                "severity": f.severity,
                "volume_cc": f.volume_cc,
                "description": f.description,
                "recommendation": f.recommendation,
            }
            for f in study.findings
        ]
        narrative = vila.report_narrative(
            str(study.id),
            patient_name=study.patient_name,
            modality=study.modality,
            body_part=study.body_part,
            description=study.description,
            findings=findings_payload,
            series_instances=_series_instances_for_study(study),
        )
        if narrative.get("narrative"):
            ctx["ai_narrative"] = narrative["narrative"]
            ctx["ai_impression"] = narrative.get("impression")
    return ctx


def render_html(study: Study) -> str:
    return _env.get_template("report.html").render(**_context(study))


def render_pdf(html: str) -> bytes | None:
    try:
        from weasyprint import HTML  # noqa: PLC0415
    except Exception as exc:  # noqa: BLE001
        logger.warning("WeasyPrint unavailable, skipping PDF: %s", exc)
        return None
    return HTML(string=html).write_pdf()


def render_docx(study: Study) -> bytes | None:
    try:
        import io  # noqa: PLC0415

        from docx import Document  # noqa: PLC0415
    except Exception as exc:  # noqa: BLE001
        logger.warning("python-docx unavailable, skipping DOCX: %s", exc)
        return None

    doc = Document()
    doc.add_heading(f"{settings.LUMIRA_APP_NAME} — AI-Assisted Imaging Report", level=0)

    doc.add_heading("Patient Information", level=1)
    doc.add_paragraph(f"Name: {study.patient_name or '—'}")
    doc.add_paragraph(f"Patient ID: {study.patient_id or '—'}")
    doc.add_paragraph(f"Sex: {study.patient_sex or '—'}    Age: {study.patient_age or '—'}")

    doc.add_heading("Study Information", level=1)
    doc.add_paragraph(f"Modality: {study.modality or '—'}    Body Part: {study.body_part or '—'}")
    doc.add_paragraph(f"Description: {study.description or '—'}")

    doc.add_heading("Findings", level=1)
    if study.findings:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, name in enumerate(["Finding", "Location", "Confidence", "Volume", "Severity"]):
            hdr[i].text = name
        for f in sorted(study.findings, key=lambda x: x.confidence, reverse=True):
            row = table.add_row().cells
            row[0].text = f.label
            row[1].text = f.location
            row[2].text = f"{int(f.confidence * 100)}%"
            row[3].text = f"{f.volume_cc:.1f} cc" if f.volume_cc else "—"
            row[4].text = f.severity.upper()
    else:
        doc.add_paragraph("No findings detected.")

    doc.add_heading("Recommendations", level=1)
    if study.findings:
        for f in study.findings:
            doc.add_paragraph(f"{f.label}: {f.recommendation}", style="List Bullet")
    else:
        doc.add_paragraph("No specific recommendations.")

    doc.add_paragraph(
        "\nGenerated by AI-assisted analysis. Not a medical diagnosis. Not for clinical use."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_and_store(study: Study) -> dict[str, str]:
    """Render all formats, store in MinIO, return {format: object_key}."""
    keys: dict[str, str] = {}
    base = f"studies/{study.id}/reports"

    html = render_html(study)
    html_key = f"{base}/report.html"
    storage.put_object(html_key, html.encode("utf-8"), "text/html")
    keys["html"] = html_key

    pdf = render_pdf(html)
    if pdf:
        pdf_key = f"{base}/report.pdf"
        storage.put_object(pdf_key, pdf, "application/pdf")
        keys["pdf"] = pdf_key

    docx = render_docx(study)
    if docx:
        docx_key = f"{base}/report.docx"
        storage.put_object(
            docx_key,
            docx,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        keys["docx"] = docx_key

    return keys
